# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from dataclasses import asdict, dataclass
from typing import Any, Optional

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse, Response
from pydantic import BaseModel

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font
except Exception:
    Workbook = None
    load_workbook = None


SUPPORTED_EXTS = {".txt", ".pdf", ".docx", ".doc"}
STATE_FILE = "grading_state_web.json"
CACHE_DIR = ".grader_preview_cache"
_EXPORT_CACHE: dict[str, dict[str, Any]] = {}
_EXPORT_TEMP_DIR = os.path.join(tempfile.gettempdir(), "homework_grader_exports")
_EXPORT_TTL = 600  # 10 minutes


def _purge_stale_exports() -> None:
    now = time.time()
    stale = [t for t, v in _EXPORT_CACHE.items() if now - v["created_at"] > _EXPORT_TTL]
    for token in stale:
        try:
            os.remove(_EXPORT_CACHE[token]["path"])
        except Exception:
            pass
        del _EXPORT_CACHE[token]


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    try:
        text = str(value)
    except Exception:
        text = repr(value)
    return "".join(ch for ch in text if not 0xD800 <= ord(ch) <= 0xDFFF)


def safe_basename(path_or_name: Any) -> str:
    text = clean_text(path_or_name).replace("\\", "/")
    return text.rstrip("/").split("/")[-1]


def safe_suffix(path_or_name: Any) -> str:
    return os.path.splitext(safe_basename(path_or_name))[1].lower()


def safe_stem(path_or_name: Any) -> str:
    return os.path.splitext(safe_basename(path_or_name))[0]


def normalize_path_key(path_or_name: Any) -> str:
    return clean_text(path_or_name).replace("\\", "/").strip()


def extract_student_id(path_or_name: Any) -> str:
    """从相对路径中提取学号：遍历每一层目录和文件名，返回最长的数字匹配。

    规则（按优先级）：
    1. 每层路径片段的开头连续数字（>=1位）
    2. 每层路径片段中 >=4 位的连续数字（兼容"张三_2024010001"）
    3. 取所有候选中长度最长的
    4. 兜底：文件名的 stem
    """
    text = clean_text(path_or_name).replace("\\", "/")
    best = ""
    for part in text.split("/"):
        stem = os.path.splitext(part)[0]
        # 前缀数字
        prefix = re.match(r"^\s*(\d+)", stem)
        if prefix:
            candidate = prefix.group(1)
            if len(candidate) > len(best):
                best = candidate
        # 非前缀但 >=4 位的数字串
        for m in re.finditer(r"\d{4,}", stem):
            if len(m.group()) > len(best):
                best = m.group()
    if best:
        return best
    return safe_stem(safe_basename(text))


def normalize_student_id(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return clean_text(value).strip()
    text = clean_text(value).strip()
    if re.fullmatch(r"\d+\.0", text):
        return text[:-2]
    sci = re.fullmatch(r"(\d+(?:\.\d+)?)[eE]\+?(\d+)", text)
    if sci:
        try:
            return str(int(float(text)))
        except Exception:
            pass
    m = re.search(r"\d+", text)
    return m.group(0) if m else text


def make_file_key(path: str, folder: str, recursive: bool) -> str:
    if recursive:
        try:
            rel = os.path.relpath(path, folder)
            if not rel.startswith(".."):
                return normalize_path_key(rel)
        except Exception:
            pass
    return normalize_path_key(safe_basename(path))


def display_name(path: str, folder: str, recursive: bool) -> str:
    if recursive:
        try:
            rel = os.path.relpath(path, folder)
            if not rel.startswith(".."):
                return normalize_path_key(rel)
        except Exception:
            pass
    return safe_basename(path)


def read_txt_file(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            with open(path, "r", encoding=enc) as f:
                return clean_text(f.read())
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            return f"无法读取 TXT 文件：{exc}"
    try:
        with open(path, "r", errors="ignore") as f:
            return clean_text(f.read())
    except Exception as exc:
        return f"无法读取 TXT 文件：{exc}"


def read_docx_text(path: str) -> str:
    if Document is None:
        return "缺少 python-docx 依赖，无法提取 Word 文本。"
    if safe_suffix(path) != ".docx":
        return "该文件不是 .docx 格式，无法用 python-docx 提取文本。"
    try:
        if not zipfile.is_zipfile(path):
            return "该 .docx 文件不是有效的 Office Open XML 压缩包。"
        doc = Document(path)
        parts: list[str] = []
        for p in doc.paragraphs:
            text = clean_text(p.text).strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append("\t".join(clean_text(cell.text).strip() for cell in row.cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts).strip() or "Word 文档中未提取到文字。"
    except BaseException as exc:
        return f"鏃犳硶瀹夊叏鎻愬彇 Word 绾枃鏈€傚師鍥狅細{type(exc).__name__}: {exc}"


def file_cache_name(path: str, suffix: str = ".pdf") -> str:
    try:
        stat = os.stat(path)
        key = f"{os.path.abspath(path)}::{stat.st_mtime_ns}::{stat.st_size}".encode("utf-8", "ignore")
    except Exception:
        key = os.path.abspath(path).encode("utf-8", "ignore")
    return hashlib.md5(key).hexdigest() + suffix


def ensure_dir(path: str) -> str:
    if os.path.exists(path) and not os.path.isdir(path):
        path += "_dir"
    os.makedirs(path, exist_ok=True)
    return path


def find_libreoffice() -> Optional[str]:
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    if sys.platform.startswith("win"):
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if os.path.exists(candidate):
                return candidate
    if sys.platform == "darwin":
        candidate = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if os.path.exists(candidate):
            return candidate
    return None


def convert_with_libreoffice(src: str, out_dir: str) -> str:
    exe = find_libreoffice()
    if not exe:
        raise RuntimeError("未找到 LibreOffice/soffice。")
    out_dir = ensure_dir(out_dir)
    cmd = [exe, "--headless", "--convert-to", "pdf", "--outdir", out_dir, src]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=180)
    pdf_path = os.path.join(out_dir, safe_stem(src) + ".pdf")
    if result.returncode != 0 or not os.path.exists(pdf_path):
        raise RuntimeError((result.stderr or result.stdout or "LibreOffice 杞崲澶辫触").strip())
    return pdf_path


def convert_with_word_com(src: str, out_pdf: str) -> str:
    if not sys.platform.startswith("win"):
        raise RuntimeError("Microsoft Word COM 仅支持 Windows。")
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError("未安装 pywin32，无法调用 Microsoft Word。") from exc

    ensure_dir(os.path.dirname(out_pdf))
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(src))
        doc.ExportAsFixedFormat(os.path.abspath(out_pdf), 17)
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        finally:
            try:
                if word is not None:
                    word.Quit()
            finally:
                pythoncom.CoUninitialize()
    if not os.path.exists(out_pdf):
        raise RuntimeError("Microsoft Word 未生成 PDF 文件。")
    return out_pdf


def office_to_pdf(src: str, cache_dir: str) -> str:
    cache_dir = ensure_dir(cache_dir)
    cached = os.path.join(cache_dir, file_cache_name(src, ".pdf"))
    try:
        if os.path.exists(cached) and os.path.getsize(cached) > 0:
            return cached
    except Exception:
        pass

    temp_dir = ensure_dir(os.path.join(cache_dir, "tmp_convert"))
    errors: list[str] = []
    try:
        pdf = convert_with_libreoffice(src, temp_dir)
        shutil.move(pdf, cached)
        return cached
    except BaseException as exc:
        errors.append(f"LibreOffice: {type(exc).__name__}: {exc}")
    try:
        return convert_with_word_com(src, cached)
    except BaseException as exc:
        errors.append(f"Microsoft Word: {type(exc).__name__}: {exc}")
    raise RuntimeError("\n".join(errors))


def read_roster_file(path: str) -> dict[str, str]:
    suffix = safe_suffix(path)
    rows: list[list[Any]] = []
    if suffix == ".xlsx":
        if load_workbook is None:
            raise RuntimeError("?? openpyxl????? Excel ???")
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        for row in ws.iter_rows(values_only=True):
            rows.append(list(row))
        wb.close()
    elif suffix in {".csv", ".txt"}:
        content = None
        for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
            try:
                with open(path, "r", encoding=enc, newline="") as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        if content is None:
            with open(path, "r", errors="ignore", newline="") as f:
                content = f.read()
        try:
            dialect = csv.Sniffer().sniff(content[:4096], delimiters=",	;?")
        except Exception:
            dialect = csv.excel
        rows = [row for row in csv.reader(content.splitlines(), dialect)]
    else:
        raise RuntimeError("??????? .xlsx?.csv?.txt?")

    rows = [r for r in rows if any(clean_text(c).strip() for c in r)]
    if not rows:
        raise RuntimeError("???????")

    sid_col = None
    name_col = None
    header_idx = 0
    for i in range(min(len(rows), 10)):
        headers = [clean_text(c).strip().lower() for c in rows[i]]
        for j, h in enumerate(headers):
            if sid_col is None and ("??" in h or "????" in h or h in {"id", "student id", "number", "no"}):
                sid_col = j
            if name_col is None and ("??" in h or "??" in h or h in {"name", "student name"}):
                name_col = j
        if sid_col is not None and name_col is not None:
            header_idx = i + 1
            break
        sid_col = None
        name_col = None
    if sid_col is None or name_col is None:
        sid_col = 0
        name_col = 1 if len(rows[0]) > 1 else 0
        header_idx = 0

    roster: dict[str, str] = {}
    for row in rows[header_idx:]:
        sid = normalize_student_id(row[sid_col] if sid_col < len(row) else "")
        name = clean_text(row[name_col] if name_col < len(row) else "").strip()
        if not sid or sid.lower() in {"??", "studentid", "id"}:
            continue
        if name in {"??", "??"}:
            name = ""
        roster[sid] = name
    if not roster:
        raise RuntimeError("???????????????")
    return roster

@dataclass
class GradeRecord:
    filename: str = ""
    score: str = ""
    comment: str = ""
    status: str = "未批改"


class ScanRequest(BaseModel):
    folder: str
    recursive: bool = True


class ImportRosterRequest(BaseModel):
    path: str
    existing: dict[str, str] = {}
    mode: str = "replace"


class StatePayload(BaseModel):
    folder: str
    recursive: bool = True
    roster: dict[str, str] = {}
    records: dict[str, GradeRecord] = {}


class ExportRequest(StatePayload):
    format: str = "csv"


class ExportPrepareRequest(StatePayload):
    kind: str = "csv"  # "csv" or "xlsx"


class PreviewRequest(BaseModel):
    path: str
    folder: str = ""
    prefer_pdf: bool = True


app = FastAPI(title="Homework Grader API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def add_desktop_cors_headers(request, call_next):
    if request.method == "OPTIONS":
        response = Response()
    else:
        response = await call_next(request)
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,PUT,PATCH,DELETE,OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "*"
    return response


def assert_existing_folder(folder: str) -> str:
    folder = clean_text(folder)
    if not folder or not os.path.isdir(folder):
        raise HTTPException(status_code=400, detail="作业文件夹不存在。")
    return folder


def assert_existing_file(path: str) -> str:
    path = clean_text(path)
    if not path or not os.path.isfile(path):
        raise HTTPException(status_code=400, detail="文件不存在。")
    return path


def scan_folder(folder: str, recursive: bool) -> list[dict[str, Any]]:
    folder = assert_existing_folder(folder)
    files: list[str] = []
    try:
        if recursive:
            for root, _, names in os.walk(folder):
                if safe_basename(root) == CACHE_DIR:
                    continue
                for name in names:
                    full = os.path.join(root, name)
                    if os.path.isfile(full) and safe_suffix(name) in SUPPORTED_EXTS:
                        files.append(full)
        else:
            for name in os.listdir(folder):
                full = os.path.join(folder, name)
                if os.path.isfile(full) and safe_suffix(name) in SUPPORTED_EXTS:
                    files.append(full)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"璇诲彇鏂囦欢澶瑰け璐ワ細{exc}") from exc

    files.sort(key=lambda p: display_name(p, folder, recursive))
    submissions = []
    for path in files:
        display = display_name(path, folder, recursive)
        try:
            stat = os.stat(path)
            mtime = stat.st_mtime
            size = stat.st_size
        except Exception:
            mtime = 0
            size = 0
        safe_path = clean_text(path)
        safe_display = clean_text(display)
        safe_name = clean_text(safe_basename(path))
        submissions.append(
            {
                "key": clean_text(make_file_key(path, folder, recursive)),
                "path": safe_path,
                "displayPath": safe_display,
                "filename": safe_name,
                "studentId": clean_text(extract_student_id(display)),
                "ext": clean_text(safe_suffix(path)),
                "size": size,
                "mtime": mtime,
            }
        )
    return submissions


def pick_best(keys: list[str], records: dict[str, GradeRecord], submissions_by_key: dict[str, dict[str, Any]]) -> str:
    def score_rank(value: Any) -> float:
        text = clean_text(value).strip()
        if not text:
            return -1.0
        try:
            return float(text)
        except Exception:
            return 0.0

    def rank(key: str):
        record = records.get(key, GradeRecord(filename=submissions_by_key[key]["displayPath"]))
        sub = submissions_by_key.get(key, {})
        return (
            1 if record.status and record.status != "未批改" else 0,
            1 if record.score else 0,
            1 if record.comment else 0,
            score_rank(record.score),
            float(sub.get("mtime", 0)),
        )

    return sorted(keys, key=rank)[-1] if keys else ""


def build_export_rows(payload: StatePayload) -> list[list[str]]:
    submissions = scan_folder(payload.folder, payload.recursive)
    submissions_by_key = {item["key"]: item for item in submissions}
    grouped: dict[str, list[str]] = {}
    sid_order: list[str] = []
    for item in submissions:
        sid = item["studentId"]
        if sid not in grouped:
            grouped[sid] = []
            sid_order.append(sid)
        grouped[sid].append(item["key"])

    rows: list[list[str]] = []
    exported_sid: set[str] = set()
    if payload.roster:
        for sid, student_name in payload.roster.items():
            keys = grouped.get(sid, [])
            if not keys:
                rows.append(["", sid, clean_text(student_name), "", "未交", ""])
            else:
                key = pick_best(keys, payload.records, submissions_by_key)
                sub = submissions_by_key[key]
                record = payload.records.get(key, GradeRecord(filename=sub["displayPath"]))
                rows.append([sub["displayPath"], sid, clean_text(student_name), record.score, record.status, record.comment])
            exported_sid.add(sid)
    for sid in sid_order:
        if sid in exported_sid:
            continue
        key = pick_best(grouped.get(sid, []), payload.records, submissions_by_key)
        if not key:
            continue
        sub = submissions_by_key[key]
        record = payload.records.get(key, GradeRecord(filename=sub["displayPath"]))
        rows.append([sub["displayPath"], sid, "", record.score, record.status, record.comment])
    return rows


@app.get("/api/health")
def health():
    return {"ok": True}


@app.post("/api/folders/scan")
def scan(req: ScanRequest):
    submissions = scan_folder(req.folder, req.recursive)
    return {"folder": clean_text(req.folder), "recursive": req.recursive, "submissions": submissions}


@app.post("/api/folders/pick")
def pick_folder():
    if sys.platform.startswith("win"):
        ps_script = r"""
Add-Type -AssemblyName System.Windows.Forms
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
$dialog = New-Object System.Windows.Forms.FolderBrowserDialog
$dialog.Description = 'Select homework folder'
$dialog.ShowNewFolderButton = $false
if ($dialog.ShowDialog() -eq [System.Windows.Forms.DialogResult]::OK) {
  Write-Output $dialog.SelectedPath
}
"""
        try:
            ps_exe = (
                r"C:\Windows\System32\WindowsPowerShell\v1.0\powershell.exe"
                if os.path.exists(r"C:\Windows\System32\WindowsPowerShell\v1.0\powershell.exe")
                else "powershell"
            )
            result = subprocess.run(
                [ps_exe, "-NoProfile", "-STA", "-Command", ps_script],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding="utf-8",
                errors="ignore",
                timeout=300,
            )
            if result.returncode != 0:
                raise RuntimeError((result.stderr or result.stdout or "Folder picker failed.").strip())
            return {"folder": clean_text(result.stdout).strip()}
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"无法打开文件夹选择器：{exc}") from exc

    try:
        import tkinter as tk
        from tkinter import filedialog
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"无法打开文件夹选择器：{exc}") from exc

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    try:
        folder = filedialog.askdirectory(title="选择作业文件夹")
    finally:
        root.destroy()
    return {"folder": folder or ""}


@app.post("/api/roster/import")
def import_roster(req: ImportRosterRequest):
    roster = read_roster_file(assert_existing_file(req.path))
    if req.mode == "merge":
        merged = dict(req.existing)
        merged.update(roster)
        roster = merged
    return {"roster": roster, "count": len(roster)}


@app.post("/api/roster/upload")
async def upload_roster(file: UploadFile = File(...)):
    suffix = safe_suffix(file.filename or "")
    if suffix not in {".xlsx", ".csv", ".txt"}:
        raise HTTPException(status_code=400, detail="名单文件仅支持 .xlsx、.csv、.txt。")
    fd, temp_path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    try:
        with open(temp_path, "wb") as f:
            f.write(await file.read())
        roster = read_roster_file(temp_path)
        return {"roster": roster, "count": len(roster)}
    finally:
        try:
            os.remove(temp_path)
        except Exception:
            pass


@app.post("/api/state/save")
def save_state(payload: StatePayload):
    folder = assert_existing_folder(payload.folder)
    state_path = os.path.join(folder, STATE_FILE)
    data = payload.model_dump()
    with open(state_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return {"saved": True, "path": state_path}


@app.get("/api/state/load")
def load_state(folder: str):
    folder = assert_existing_folder(folder)
    state_path = os.path.join(folder, STATE_FILE)
    if not os.path.exists(state_path):
        return {"exists": False}
    with open(state_path, "r", encoding="utf-8") as f:
        return {"exists": True, "state": json.load(f)}


@app.post("/api/export/prepare")
def export_prepare(payload: ExportPrepareRequest):
    """生成导出文件并返回一次性下载 token。"""
    _purge_stale_exports()

    rows = build_export_rows(payload)
    os.makedirs(_EXPORT_TEMP_DIR, exist_ok=True)

    token = hashlib.md5(os.urandom(16)).hexdigest()
    dirname = os.path.basename(payload.folder.rstrip("/\\")) or "成绩表"
    filename = f"{dirname}成绩表.csv" if payload.kind == "csv" else f"{dirname}成绩表.xlsx"
    out_path = os.path.join(_EXPORT_TEMP_DIR, f"{token}_{filename}")

    if payload.kind == "csv":
        with open(out_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["文件名", "学号", "姓名", "分数", "状态", "评语"])
            writer.writerows(rows)
    else:
        if Workbook is None:
            raise HTTPException(status_code=500, detail="缺少 openpyxl。")
        wb = Workbook()
        ws = wb.active
        ws.title = "成绩表"
        headers = ["文件名", "学号", "姓名", "分数", "状态", "评语"]
        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")
        for row in rows:
            ws.append(row)
        for col, width in {"A": 52, "B": 18, "C": 14, "D": 12, "E": 12, "F": 48}.items():
            ws.column_dimensions[col].width = width
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        wb.save(out_path)

    _EXPORT_CACHE[token] = {"path": out_path, "filename": filename, "created_at": time.time()}
    return {"token": token, "filename": filename}


@app.get("/api/export/download/{token}")
def export_download(token: str):
    """通过 token 下载导出文件（一次性）。"""
    entry = _EXPORT_CACHE.pop(token, None)
    if not entry:
        raise HTTPException(status_code=404, detail="下载链接已过期，请重新导出。")
    media = "text/csv" if entry["filename"].endswith(".csv") else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return FileResponse(
        entry["path"],
        filename=entry["filename"],
        media_type=media,
    )


@app.post("/api/export/csv")
def export_csv(payload: ExportRequest):
    rows = build_export_rows(payload)
    out_path = os.path.join(assert_existing_folder(payload.folder), "成绩表.csv")
    with open(out_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["文件名", "学号", "姓名", "分数", "状态", "评语"])
        writer.writerows(rows)
    return FileResponse(out_path, filename="成绩表.csv", media_type="text/csv")


@app.post("/api/export/xlsx")
def export_xlsx(payload: ExportRequest):
    if Workbook is None:
        raise HTTPException(status_code=500, detail="缺少 openpyxl。")
    rows = build_export_rows(payload)
    out_path = os.path.join(assert_existing_folder(payload.folder), "成绩表.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "成绩表"
    headers = ["文件名", "学号", "姓名", "分数", "状态", "评语"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
    for row in rows:
        ws.append(row)
    for col, width in {"A": 52, "B": 18, "C": 14, "D": 12, "E": 12, "F": 48}.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(out_path)
    return FileResponse(out_path, filename="成绩表.xlsx", media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.post("/api/preview/text")
def preview_text(req: PreviewRequest):
    path = assert_existing_file(req.path)
    suffix = safe_suffix(path)
    if suffix == ".txt":
        return {"kind": "text", "text": read_txt_file(path)}
    if suffix == ".docx":
        return {"kind": "text", "text": read_docx_text(path)}
    if suffix == ".doc":
        return {"kind": "text", "text": "旧版 .doc 格式暂不能直接提取文本，请使用 PDF 预览或外部打开。"}
    if suffix == ".pdf":
        return {"kind": "text", "text": "PDF 请使用 PDF 预览。"}
    raise HTTPException(status_code=400, detail="不支持的预览格式。")


@app.post("/api/preview/pdf")
def preview_pdf(req: PreviewRequest):
    path = assert_existing_file(req.path)
    suffix = safe_suffix(path)
    if suffix == ".pdf":
        return FileResponse(path, media_type="application/pdf")
    if suffix in {".doc", ".docx"}:
        folder = assert_existing_folder(req.folder) if req.folder else os.path.dirname(path)
        cache_dir = os.path.join(folder, CACHE_DIR)
        try:
            pdf = office_to_pdf(path, cache_dir)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"无法转换为 PDF：{exc}") from exc
        return FileResponse(pdf, media_type="application/pdf")
    raise HTTPException(status_code=400, detail="该文件不能生成 PDF 预览。")


@app.post("/api/cache/clear")
def clear_cache(req: ScanRequest):
    folder = assert_existing_folder(req.folder)
    cache_dir = os.path.abspath(os.path.join(folder, CACHE_DIR))
    expected = os.path.abspath(os.path.join(folder, CACHE_DIR))
    if cache_dir == expected and os.path.isdir(cache_dir):
        shutil.rmtree(cache_dir, ignore_errors=True)
    return {"cleared": True}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8765, log_config=None, access_log=False)

